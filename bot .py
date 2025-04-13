
#!/usr/bin/env python3
# -*- coding: utf-8 -*-

# Установка зависимостей
# Закомментировано, т.к. в Colab лучше выполнять в отдельной ячейке
!pip install python-docx docxcompose beautifulsoup4 ebooklib aiogram aiofiles nest_asyncio

import os
import re
import time
import docx
import aiogram
from docx import Document
from docxcompose.composer import Composer
from bs4 import BeautifulSoup
import ebooklib
from ebooklib import epub
from aiogram import Bot, Router, types, F, Dispatcher
from aiogram.types import Message, FSInputFile, ReplyKeyboardRemove, CallbackQuery
from aiogram.filters import Command
from aiogram.utils.keyboard import ReplyKeyboardBuilder, InlineKeyboardBuilder
from aiogram.utils import markdown as md
import aiofiles
import asyncio
import nest_asyncio
import concurrent.futures
from aiogram.fsm.context import FSMContext
from aiogram.fsm.state import State
from aiogram.fsm.state import StatesGroup
from aiogram.fsm.storage.memory import MemoryStorage
from functools import partial
from collections import deque
from datetime import datetime, timezone, timedelta
from aiogram.exceptions import TelegramBadRequest # Для обработки ошибок удаления/редактирования

nest_asyncio.apply()
from aiogram.types import InlineKeyboardMarkup, InlineKeyboardButton

# Создаем пул потоков для выполнения CPU-bound задач
thread_pool = concurrent.futures.ThreadPoolExecutor(max_workers=1)

# --- Вспомогательная функция для удаления сообщений ---
async def delete_message_after_delay(message: types.Message, delay: int):
    """Удаляет сообщение после указанной задержки."""
    await asyncio.sleep(delay)
    try:
        await message.delete()
    except TelegramBadRequest: # Игнорируем ошибки, если сообщение уже удалено
        pass
    except Exception as e:
        print(f"Не удалось удалить сообщение {message.message_id}: {e}")

# --- Вспомогательная функция для отправки и последующего удаления ---
async def send_then_delete(chat_id: int, text: str, delay: int = 5, **kwargs):
    """Отправляет сообщение и планирует его удаление."""
    try:
        sent_message = await bot.send_message(chat_id, text, **kwargs)
        asyncio.create_task(delete_message_after_delay(sent_message, delay))
    except Exception as e:
        print(f"Ошибка при отправке или планировании удаления: {e}")

class UserLimits:
    def __init__(self, max_files, max_size):
        self.user_data = {}  # {user_id: {'files_today': int}}
        self.last_global_reset = self._get_last_utc_midnight()
        self.user_locks = {} # Словарь для хранения блокировок пользователей
        self.max_files = max_files
        self.max_size = max_size

    def _get_last_utc_midnight(self):
        """Возвращает последнюю полночь по UTC."""
        now = datetime.now(timezone.utc)
        return now.replace(hour=0, minute=0, second=0, microsecond=0)

    def get_lock(self, user_id):
        """Получает или создает блокировку для пользователя."""
        if user_id not in self.user_locks:
            self.user_locks[user_id] = asyncio.Lock()
        return self.user_locks[user_id]

    def check_limits(self, user_id, file_size):
        """Проверяет лимиты и сбрасывает их в 00:00 UTC."""
        now = datetime.now(timezone.utc)

        # Если наступил новый день (00:00 UTC), сбрасываем счетчики у всех
        if now >= self.last_global_reset + timedelta(days=1): # Используем >= для надежности
            print("Сброс суточных лимитов...")
            self.user_data.clear()  # Обнуляем данные всех пользователей
            self.last_global_reset = self._get_last_utc_midnight()

        # Инициализируем данные пользователя, если их нет
        if user_id not in self.user_data:
            self.user_data[user_id] = {'files_today': 0}

        # Проверяем лимиты
        if file_size > self.max_size * 1024 * 1024:  # Допустимый размер файла
            return False, f"❌ Размер файла превышает {self.max_size} MB."

        if self.user_data[user_id]['files_today'] >= self.max_files: # Используем >=
            time_left = (self.last_global_reset + timedelta(days=1)) - now
            hours_left = int(time_left.total_seconds() // 3600)
            minutes_left = int((time_left.total_seconds() % 3600) // 60)
            return False, f"❌ Лимит исчерпан ({self.user_data[user_id]['files_today']}/{self.max_files}). Сброс через {hours_left} ч. {minutes_left} мин. (в 00:00 UTC)."

        return True, ""

    def increment_counter(self, user_id):
        """Увеличивает счетчик файлов пользователя."""
        if user_id not in self.user_data: # Инициализируем на всякий случай
             self.user_data[user_id] = {'files_today': 0}
        self.user_data[user_id]['files_today'] += 1

async def check_sender(message: types.Message) -> bool:
    """Проверяет отправителя. Если не пользователь, отвечает и возвращает True."""
    if message.sender_chat:
        await message.reply("Анонимные пользователи (от имени каналов/групп) не могут использовать этого бота.")
        # Не удаляем сообщение, т.к. оно может быть от админа канала
        return True # Да, это не пользователь, обработку надо прервать
    return False # Нет, это пользователь, можно продолжать

# Создаем экземпляр класса лимитов
user_limits = UserLimits(max_files=30, max_size=15)

# --- Глобальный словарь для хранения ID сообщений со статусом очереди ---
user_status_messages = {} # {user_id: message_id}

# Система очереди
class TaskQueue:
    def __init__(self, max_concurrent_tasks):
        self.queue = deque()  # Очередь задач
        self.active_tasks = {}  # Активные задачи: task_id -> task
        self.max_concurrent_tasks = max_concurrent_tasks
        self.task_counter = 0  # Счетчик задач для назначения номера очереди

    def add_task(self, user_id, chat_id, message_thread_id, is_forum, file_list, output_file_name):
        """Добавляет задачу в очередь и возвращает позицию в очереди"""
        self.task_counter += 1
        task_id = self.task_counter
        task = {
            'user_id': user_id,
            'chat_id': chat_id,
            'message_thread_id': message_thread_id,
            'is_forum': is_forum,
            'file_list': file_list,
            'output_file_name': output_file_name,
            'task_id': task_id,
            'time_added': time.time()
        }
        self.queue.append(task)
        return len(self.queue) # Возвращаем текущую позицию в очереди

    def get_next_task(self):
        """Получить следующую задачу из очереди"""
        if not self.queue:
            return None
        task = self.queue.popleft()
        self.active_tasks[task['task_id']] = task
        return task

    async def complete_task(self, task_id):
        """Пометить задачу как завершенную и обновить статус, если нужно."""
        if task_id in self.active_tasks:
            task = self.active_tasks.pop(task_id) # Удаляем и получаем задачу
            user_id = task['user_id']

            # Обновляем сообщение со статусом очереди пользователя, если оно есть
            if user_id in user_status_messages:
                message_id = user_status_messages[user_id]
                chat_id = task['chat_id'] # Берем chat_id из задачи
                try:
                    # Получаем текущую страницу из старой клавиатуры, если она есть
                    current_page = 1
                    try:
                        # Пытаемся получить сообщение, чтобы проверить его клавиатуру
                        msg = await bot.edit_message_reply_markup(chat_id, message_id, reply_markup=None) # Временное удаление клавиатуры для проверки
                        if msg and msg.reply_markup:
                           # Ищем кнопку пагинации, чтобы извлечь страницу
                           for row in msg.reply_markup.inline_keyboard:
                                for button in row:
                                    if button.callback_data.startswith("page:"):
                                        # Простой способ - взять страницу из первой кнопки пагинации
                                        current_page = int(button.callback_data.split(":")[1])
                                        break # Достаточно одной кнопки
                                if current_page > 1: break # Выход из внешнего цикла
                    except Exception: # Если сообщение не найдено или нет клавиатуры
                        pass # Остаемся на странице 1

                    text, keyboard = await build_task_status(user_id, page=current_page) # Перестраиваем статус для ТЕКУЩЕЙ страницы
                    if text: # Если есть что показывать
                         await bot.edit_message_text(text, chat_id, message_id, reply_markup=keyboard)
                    else: # Если задач не осталось, удаляем сообщение со статусом
                        await bot.delete_message(chat_id, message_id)
                        del user_status_messages[user_id]
                except TelegramBadRequest as e:
                    print(f"Не удалось обновить статус для user {user_id} (сообщение {message_id}): {e}")
                    if user_id in user_status_messages:
                         del user_status_messages[user_id] # Удаляем ID, если сообщение недоступно
                except Exception as e:
                    print(f"Другая ошибка при обновлении статуса для user {user_id}: {e}")
                    if user_id in user_status_messages:
                        del user_status_messages[user_id]

    def get_user_tasks(self, user_id):
        """Получить список всех задач пользователя (в очереди и активных)"""
        tasks = []
        # Ищем в активных задачах
        for task in self.active_tasks.values():
            if task['user_id'] == user_id:
                tasks.append(task)
        # Ищем в очереди
        queue_tasks = [task for task in self.queue if task['user_id'] == user_id]
        # Сортируем задачи: сначала активные, потом по времени добавления в очередь
        tasks.sort(key=lambda x: x['time_added'])
        queue_tasks.sort(key=lambda x: x['time_added'])
        return tasks + queue_tasks # Возвращаем сначала активные, потом в очереди

    def can_process_now(self):
        """Проверка, можно ли обработать следующую задачу из очереди"""
        return len(self.active_tasks) < self.max_concurrent_tasks and self.queue

# Создаем очередь задач
task_queue = TaskQueue(max_concurrent_tasks=1)  # Пока оставим 1

# Декоратор для измерения времени выполнения функции
def timer(func):
    async def wrapper(*args, **kwargs):
        start_time = time.time()
        result = await func(*args, **kwargs)
        elapsed = time.time() - start_time
        print(f"[PROFILING] Функция {func.__name__} выполнилась за {elapsed:.2f} секунд")
        return result
    return wrapper

# Замените токен на свой
API_TOKEN = '7690726183:AAEt9jgeOxa6dJSYn0SyDyLT6v0aaDibTa0' # Используйте переменную окружения для токена в реальных проектах

bot = Bot(token=API_TOKEN)
router = Router()

# ===================== Неблокирующие функции конвертации =====================

# Функция-обертка для выполнения блокирующих операций в отдельном потоке
async def run_in_threadpool(func, *args, **kwargs):
    loop = asyncio.get_running_loop()
    func_partial = partial(func, *args, **kwargs)
    return await loop.run_in_executor(thread_pool, func_partial)

# Неблокирующие версии функций конвертации
async def convert_epub_to_docx(epub_file, docx_file):
    def _convert():
        book = epub.read_epub(epub_file)
        document = Document()
        for item in book.get_items():
            if item.get_type() == ebooklib.ITEM_DOCUMENT:
                soup = BeautifulSoup(item.content, 'html.parser')
                # --- Улучшенная обработка EPUB ---
                body = soup.find('body')
                if not body: continue # Пропускаем, если нет тега body

                for element in body.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p']):
                    text = element.get_text().strip()
                    if not text: continue # Пропускаем пустые теги

                    level = 0
                    if element.name.startswith('h'):
                        try:
                            level = int(element.name[1])
                        except ValueError:
                            level = 1 # По умолчанию для h тегов без цифры
                    if level > 0:
                         # Ограничиваем уровень заголовка для docx (0-9, но лучше меньше)
                        document.add_heading(text, level=min(level, 4))
                    elif element.name == 'p':
                        doc_paragraph = document.add_paragraph()
                        # Обработка вложенных тегов (простой вариант)
                        for sub in element.contents:
                            run_text = ""
                            is_bold = False
                            is_italic = False
                            if hasattr(sub, 'name'):
                                run_text = sub.get_text()
                                if sub.name in ['strong', 'b']:
                                    is_bold = True
                                elif sub.name in ['em', 'i']:
                                    is_italic = True
                            elif isinstance(sub, str): # Если это просто текст
                                run_text = sub
                            else: # Пропускаем другие типы узлов (комментарии и т.д.)
                                continue

                            run_text = run_text.strip() # Убираем лишние пробелы
                            if run_text:
                                run = doc_paragraph.add_run(run_text + ' ') # Добавляем пробел для разделения
                                run.bold = is_bold
                                run.italic = is_italic
                        # Удаляем лишний пробел в конце абзаца
                        if doc_paragraph.runs:
                            last_run = doc_paragraph.runs[-1]
                            if last_run.text.endswith(' '):
                                last_run.text = last_run.text[:-1]

        document.save(docx_file)
    return await run_in_threadpool(_convert)

async def convert_fb2_to_docx(fb2_file, docx_file):
    def _convert():
        try:
            # Пытаемся с разными кодировками, если utf-8 не сработала
            encodings_to_try = ['utf-8', 'windows-1251', 'iso-8859-5']
            content = None
            for enc in encodings_to_try:
                try:
                    with open(fb2_file, 'r', encoding=enc) as f:
                        content = f.read()
                    print(f"FB2 прочитан с кодировкой: {enc}")
                    break # Выходим, если успешно
                except UnicodeDecodeError:
                    print(f"Ошибка декодирования FB2 с {enc}")
                    continue # Пробуем следующую кодировку
                except Exception as read_err: # Ловим другие ошибки чтения
                    print(f"Ошибка чтения FB2 ({enc}): {read_err}")
                    continue

            if content is None:
                raise ValueError("Не удалось прочитать FB2 файл с доступными кодировками.")

            # Используем 'lxml-xml' для лучшей обработки XML
            soup = BeautifulSoup(content, 'lxml-xml')
            document = Document()
            body = soup.find('body') # Ищем основной контент
            if not body:
                 raise ValueError("Не найден тег <body> в FB2 файле.")

            # Проходим по основным секциям или直接 по абзацам внутри body
            elements_to_process = body.find_all(['section', 'p', 'title', 'epigraph', 'poem', 'cite', 'subtitle', 'text-author', 'table']) # Добавили основные теги

            for element in elements_to_process:
                 # Пропускаем элементы внутри описания книги (description)
                 if element.find_parent('description'):
                     continue

                 text = element.get_text().strip()
                 if not text and element.name != 'empty-line': continue # Пропускаем пустые, кроме 'empty-line'

                 # Обработка заголовков внутри секций
                 if element.name == 'section':
                     title_tag = element.find('title')
                     if title_tag:
                         title_text = title_tag.get_text().strip()
                         if title_text:
                             # Определяем уровень заголовка по вложенности секций (приблизительно)
                             level = len(element.find_parents('section')) + 1
                             document.add_heading(title_text, level=min(level, 4))
                     # Обрабатываем абзацы внутри секции (если они не были обработаны ранее)
                     # Этот блок может быть избыточен, если p обрабатываются отдельно
                 elif element.name == 'title' and not element.find_parent('section'):
                     # Заголовок верхнего уровня (вне секций)
                     document.add_heading(text, level=1)
                 elif element.name == 'subtitle':
                     document.add_heading(text, level=max(len(element.find_parents('section')) + 1, 2)) # Подзаголовок
                 elif element.name == 'p':
                     doc_paragraph = document.add_paragraph()
                     # Обработка вложенных тегов (strong, emphasis)
                     for sub in element.contents:
                         run_text = ""
                         is_bold = False
                         is_italic = False
                         if hasattr(sub, 'name'):
                             run_text = sub.get_text()
                             if sub.name in ['strong']: is_bold = True
                             elif sub.name in ['emphasis']: is_italic = True
                             elif sub.name == 'a': # Обработка ссылок (просто текст)
                                 run_text = sub.get_text()
                             # Можно добавить обработку других тегов по необходимости
                         elif isinstance(sub, str):
                             run_text = sub
                         else: continue

                         run_text = re.sub(r'\s+', ' ', run_text).strip() # Нормализуем пробелы
                         if run_text:
                             run = doc_paragraph.add_run(run_text + ' ')
                             run.bold = is_bold
                             run.italic = is_italic

                     if doc_paragraph.runs: # Удаляем лишний пробел в конце
                        last_run = doc_paragraph.runs[-1]
                        if last_run.text.endswith(' '):
                             last_run.text = last_run.text[:-1]
                     # Если параграф все равно пустой (например, из-за <empty-line/> внутри), удаляем его
                     # (сложно сделать в python-docx, пока пропускаем)

                 # --- Добавим обработку других тегов (упрощенно) ---
                 elif element.name in ['epigraph', 'poem', 'cite', 'text-author']:
                      # Просто добавляем как абзац, можно добавить форматирование
                      doc_paragraph = document.add_paragraph(text)
                      if element.name in ['epigraph', 'cite', 'text-author']:
                          for run in doc_paragraph.runs: run.italic = True
                 elif element.name == 'empty-line':
                      document.add_paragraph() # Добавляем пустой абзац
                 # elif element.name == 'table': # Обработка таблиц (сложнее)
                 #     # ... логика парсинга tr, th, td и создания таблицы в docx ...
                 #     document.add_paragraph(f"[Таблица: {len(element.find_all('tr'))} строк]") # Заглушка

        except Exception as e:
            print(f"Ошибка конвертации FB2 {fb2_file}: {e}")
            # Создаем пустой docx или с сообщением об ошибке, чтобы процесс не падал
            document = Document()
            document.add_paragraph(f"Ошибка конвертации файла {os.path.basename(fb2_file)}: {e}")
        finally:
            # Сохраняем документ в любом случае
             if 'document' in locals():
                document.save(docx_file)
             else: # Если даже создание документа не удалось
                 # Создаем пустой файл
                 Document().save(docx_file)

    return await run_in_threadpool(_convert)

async def convert_txt_to_docx(txt_file, docx_file):
    def _convert():
        try:
            # Пробуем разные кодировки
            encodings_to_try = ['utf-8', 'windows-1251']
            text = None
            for enc in encodings_to_try:
                try:
                    with open(txt_file, 'r', encoding=enc) as f:
                        text = f.read()
                    break
                except UnicodeDecodeError:
                    continue
                except Exception as read_err:
                    print(f"Ошибка чтения TXT ({enc}): {read_err}")
                    continue
            if text is None:
                 raise ValueError("Не удалось прочитать TXT файл.")

            document = Document()
            for line in text.splitlines():
                document.add_paragraph(line)
            document.save(docx_file)
        except Exception as e:
             print(f"Ошибка конвертации TXT {txt_file}: {e}")
             document = Document()
             document.add_paragraph(f"Ошибка конвертации файла {os.path.basename(txt_file)}: {e}")
             document.save(docx_file) # Сохраняем с ошибкой

    return await run_in_threadpool(_convert)

@timer
async def process_files(file_list):
    """
    Обрабатывает список файлов, конвертируя их в формат .docx (если требуется)
    и возвращает список имен файлов в формате .docx для последующего объединения.
    """
    converted_files = []
    conversion_tasks = []

    # Создаем задачи на конвертацию для файлов НЕ .docx
    for file in file_list:
        ext = os.path.splitext(file)[1].lower()
        if ext == ".docx":
            converted_files.append(file) # Уже готов
        else:
            docx_file = os.path.splitext(file)[0] + ".docx"
            if ext == ".txt":
                conversion_tasks.append(convert_txt_to_docx(file, docx_file))
            elif ext == ".fb2":
                conversion_tasks.append(convert_fb2_to_docx(file, docx_file))
            elif ext == ".epub":
                conversion_tasks.append(convert_epub_to_docx(file, docx_file))
            # Добавляем имя будущего файла в список, чтобы сохранить порядок
            converted_files.append(docx_file) # Имя файла, который БУДЕТ создан

    # Запускаем конвертацию параллельно (внутри run_in_threadpool)
    if conversion_tasks:
        await asyncio.gather(*conversion_tasks)

    # Проверяем, все ли сконвертированные файлы существуют
    final_files = []
    for file in converted_files:
        if os.path.exists(file):
            final_files.append(file)
        else:
             # Если какой-то файл не создался после конвертации (ошибка в _convert)
             print(f"Предупреждение: Файл {file} не был создан после попытки конвертации.")
             # Можно создать пустой файл-заглушку, чтобы слияние не упало
             # Document().save(file)
             # final_files.append(file)
             # Или просто пропустить его:
             continue

    # Возвращаем список реально существующих docx файлов в исходном порядке
    # Сопоставляем исходный порядок с существующими файлами
    result_files = []
    original_docx_names = [os.path.splitext(f)[0] + ".docx" for f in file_list]
    existing_final_files_set = set(final_files)

    for original_docx_name in original_docx_names:
        if original_docx_name in existing_final_files_set:
            result_files.append(original_docx_name)

    return result_files


# ===================== Неблокирующие функции для работы с документами =====
def check_and_add_title(doc, file_name):
    """
    Проверяет первые абзацы документа на наличие заголовка (например, "Глава ...").
    Если заголовок не найден, добавляет его на основе имени файла.
    (Функция остается блокирующей, но вызывается из run_in_threadpool)
    """
    # Паттерны для поиска заголовков глав, прологов и т.д. (регистронезависимые)
    patterns = [
        r'Глава[ ]{0,4}\d{1,4}',
        r'Часть[ ]{0,4}\d{1,4}',
        r'^Пролог[ .!]*$',
        r'^Эпилог[ .!]*$',
        r'^Описание[ .!]*$',
        r'^Аннотация[ .!]*$',
        r'^Annotation[ .!]*$',
        r'^Предисловие(\s+от\s+автора)?[ .!]*$',
        r'^Послесловие[ .!]*$',
        r'^ГЛАВА\s+\d+', # Все заглавные
        r'Chapter\s+\d+', # Английский вариант
    ]
    title_found = False
    if doc.paragraphs:
        # Проверяем первые ~5 абзацев или меньше, если их меньше
        check_limit = min(5, len(doc.paragraphs))
        for i in range(check_limit):
            p = doc.paragraphs[i]
            # Проверяем текст абзаца и его стиль
            is_heading_style = p.style.name.lower().startswith('heading')
            text_matches = any(re.search(pattern, p.text.strip(), re.IGNORECASE | re.UNICODE) for pattern in patterns)

            if text_matches or is_heading_style:
                # Дополнительная проверка: не является ли это частью обычного текста?
                # Например, если абзац длинный, это вряд ли заголовок.
                # Условие можно усложнить, но пока оставим так.
                if len(p.text.strip()) < 100: # Примерная длина заголовка
                    title_found = True
                    print(f"Найден заголовок '{p.text[:30]}...' в {os.path.basename(file_name)}")
                    break # Прекращаем поиск, если нашли

        if not title_found:
            # Извлекаем имя файла без расширения и без цифр в скобках (если были добавлены)
            base_name = os.path.splitext(os.path.basename(file_name))[0]
            # Удаляем возможные суффиксы типа "(1)", "(2)" и т.д.
            base_name = re.sub(r'\(\d+\)$', '', base_name).strip()
            # Заменяем подчеркивания и дефисы на пробелы для читаемости
            title_text = base_name.replace('_', ' ').replace('-', ' ')

            print(f"Заголовок не найден в {os.path.basename(file_name)}. Добавляем: '{title_text}'")
            # Добавляем заголовок перед первым абзацем
            # Создаем новый абзац для заголовка
            title_paragraph = doc.paragraphs[0].insert_paragraph_before(title_text)
            title_paragraph.style = doc.styles['Heading 1'] # Применяем стиль заголовка
            # Можно добавить пустую строку после заголовка для отступа
            title_paragraph.insert_paragraph_before("") # Добавляет пустой абзац *перед* заголовком
            # Правильнее добавить после:
            # title_paragraph.add_run("\n") # Не сработает как новая строка в docx
            # Вместо этого вставляем еще один пустой параграф ПОСЛЕ заголовка
            # Это делается при слиянии, см. composer.append

    return doc # Возвращаем измененный или неизмененный документ

@timer
async def merge_docx(file_list, output_file_name):
    def _merge():
        if not file_list:
            print("Нет файлов для объединения.")
            # Создаем пустой документ с сообщением
            doc = Document()
            doc.add_paragraph("Ошибка: Не найдено файлов для объединения.")
            doc.save(output_file_name)
            return output_file_name

        try:
            # Открываем первый файл и проверяем/добавляем заголовок
            print(f"Открываем первый файл: {file_list[0]}")
            merged_document = Document(file_list[0])
            merged_document = check_and_add_title(merged_document, file_list[0])
            composer = Composer(merged_document)

            # Добавляем остальные файлы
            for i, file in enumerate(file_list[1:]):
                print(f"Добавляем файл {i+2}/{len(file_list)}: {file}")
                try:
                    doc_to_append = Document(file)
                    # Проверяем и добавляем название главы при необходимости
                    doc_to_append = check_and_add_title(doc_to_append, file)
                    # Добавляем разрыв страницы перед каждым новым файлом (кроме первого)
                    # Важно: добавляем разрыв в *предыдущий* документ перед append
                    # Но Composer делает это сам или управляет этим иначе.
                    # Попробуем добавить пустой абзац для визуального разделения, если заголовка не было
                    # composer.append() сам добавляет разрыв секции по умолчанию
                    composer.append(doc_to_append)
                except Exception as append_err:
                    print(f"Ошибка при добавлении файла {file}: {append_err}")
                    # Добавляем сообщение об ошибке в итоговый документ
                    composer.doc.add_paragraph(f"\n--- Ошибка: Не удалось добавить файл {os.path.basename(file)} ---")
                    composer.doc.add_page_break() # Добавляем разрыв, чтобы продолжить

            # Сохраняем итоговый документ
            print(f"Сохранение итогового файла: {output_file_name}")
            composer.save(output_file_name)
            print(f"Файлы объединены в {output_file_name}")
            return output_file_name
        except Exception as merge_err:
            print(f"Критическая ошибка при объединении файлов: {merge_err}")
            # Создаем файл с сообщением об ошибке
            doc = Document()
            doc.add_paragraph(f"Критическая ошибка при объединении файлов: {merge_err}")
            doc.save(output_file_name)
            return output_file_name

    # Объединяем обработанные файлы в отдельном потоке
    result = await run_in_threadpool(_merge)
    return result

# ===================== FSM: Состояния =====================
class MergeStates(StatesGroup):
    collecting = State()  # Состояние сбора файлов
    naming_file = State() # Состояние запроса имени файла

# ===================== Обработчики Telegram-бота =====================

# --- Таймер автоотмены сбора ---
COLLECTION_TIMEOUT = 300 # 5 минут в секундах

async def collection_timeout_task(chat_id: int, user_id: int, state: FSMContext, initial_message_id: int):
    """Задача, которая ждет таймаут и отменяет сбор, если он еще активен."""
    await asyncio.sleep(COLLECTION_TIMEOUT)
    try:
        current_state = await state.get_state()
        # Проверяем, что состояние все еще 'collecting' и для того же пользователя/чата
        # (Теоретически, state может быть переиспользован, нужна проверка user_id)
        user_data = await state.get_data()
        # Добавим проверку, что таймер запущен для правильного пользователя
        if current_state == MergeStates.collecting.state and user_data.get("_initiator_user_id") == user_id:
            print(f"Таймаут сбора для пользователя {user_id} в чате {chat_id}. Отмена...")

            # Логика отмены (похожа на /cancel)
            file_list_data = user_data.get('file_list', [])
            saved_message_ids = user_data.get('saved_message_ids', [])
            start_message_id = user_data.get('start_message_id', initial_message_id) # ID сообщения /start_merge

            # 1. Удаляем временные файлы
            for file_item in file_list_data:
                file = file_item[0]
                if os.path.exists(file):
                    try:
                        os.remove(file)
                    except OSError as e:
                        print(f"Ошибка удаления файла {file}: {e}")

            # 2. Удаляем сообщения "Файл сохранен"
            for msg_id in saved_message_ids:
                try:
                    await bot.delete_message(chat_id, msg_id)
                    await asyncio.sleep(0.1) # Небольшая пауза
                except TelegramBadRequest: pass # Игнорируем, если уже удалено
                except Exception as e: print(f"Ошибка удаления saved_msg {msg_id}: {e}")

            # 3. Удаляем сообщение /start_merge (если есть ID)
            if start_message_id:
                 try:
                     await bot.delete_message(chat_id, start_message_id)
                 except TelegramBadRequest: pass
                 except Exception as e: print(f"Ошибка удаления start_msg {start_message_id}: {e}")

            # 4. Сбрасываем состояние
            await state.clear()

            # 5. Отправляем уведомление об отмене по таймауту (и удаляем его)
            await send_then_delete(chat_id, "⏰ Сбор файлов отменен из-за неактивности.", delay=10)

    except asyncio.CancelledError:
        # Таймаут был отменен (через /end_merge или /cancel) - это нормально
        print(f"Таймер сбора для пользователя {user_id} отменен.")
    except Exception as e:
        print(f"Ошибка в задаче таймаута: {e}")

@router.message(Command("start_merge"))
async def start_merge(message: Message, state: FSMContext):
    """
    Команда для начала сбора файлов.
    """
    if await check_sender(message):
        await message.delete() # Удаляем команду не-пользователя
        return

    user_id = message.from_user.id
    chat_id = message.chat.id
    current_state = await state.get_state()

    if current_state == MergeStates.collecting.state:
        # Проверяем, не для этого ли пользователя уже идет сбор
        user_data = await state.get_data()
        if user_data.get("_initiator_user_id") == user_id:
             await send_then_delete(message.chat.id, "Сбор файлов уже запущен. Отправляйте файлы или используйте /end_merge /cancel.", delay=5)
             await message.delete() # Удаляем повторную команду
             return
        else:
             # Если состояние есть, но от другого пользователя (маловероятно с MemoryStorage, но все же)
             await send_then_delete(message.chat.id, "В данный момент другой пользователь собирает файлы. Пожалуйста, подождите.", delay=10)
             await message.delete()
             return

    # Очищаем состояние на всякий случай перед началом нового сбора
    await state.clear()
    await state.set_state(MergeStates.collecting)
    # Сохраняем ID пользователя, начавшего сбор, и ID его сообщения /start_merge
    await state.update_data(
        file_list=[],
        saved_message_ids=[], # Список для ID сообщений "Файл сохранен"
        _initiator_user_id=user_id, # Для проверки в таймауте
        start_message_id=message.message_id # Сохраняем ID команды /start_merge
    )

    # Запускаем таймер автоотмены
    timeout_handle = asyncio.create_task(
        collection_timeout_task(chat_id, user_id, state, message.message_id)
    )
    await state.update_data(timeout_task=timeout_handle) # Сохраняем хэндл задачи таймера

    await message.answer("✅ Сбор файлов начат! Отправляйте файлы (docx, fb2, epub, txt).\n"
                         f"Используйте /end_merge для завершения или /cancel для отмены.\n"
                         f"⚠️ Сбор будет автоматически отменен через {int(COLLECTION_TIMEOUT/60)} минут бездействия.")
    # Не удаляем команду /start_merge сразу, т.к. ответ привязан к ней
    # Она удалится при таймауте, /cancel или /end_merge


# --- Функция для построения клавиатуры статуса с пагинацией ---
ITEMS_PER_PAGE = 4 # 2 кнопки в ряду * 2 ряда

async def build_task_status(user_id, page=1):
    """Строит текст и клавиатуру для статуса очереди с пагинацией."""
    user_tasks = task_queue.get_user_tasks(user_id) # Получаем задачи пользователя (активные + в очереди)

    if not user_tasks:
        total_tasks_in_queue = len(task_queue.queue)
        active_tasks_count = len(task_queue.active_tasks)
        text = (f"У вас нет активных задач или задач в очереди.\n"
                f"Статус системы: {active_tasks_count}/{task_queue.max_concurrent_tasks} активных, {total_tasks_in_queue} в очереди.")
        return text, None

    # Пагинация
    start_index = (page - 1) * ITEMS_PER_PAGE
    end_index = start_index + ITEMS_PER_PAGE
    tasks_on_page = user_tasks[start_index:end_index]

    tasks_info = []
    builder = InlineKeyboardBuilder()

    for i, task in enumerate(tasks_on_page):
        task_id = task['task_id']
        position_in_user_list = start_index + i + 1

        # Статус задачи
        if task_id in task_queue.active_tasks:
            status = "⚙️ Выполняется"
            queue_pos_str = ""
        else:
            # Ищем глобальную позицию в очереди
            global_pos = -1
            for idx, queued_task in enumerate(task_queue.queue):
                if queued_task['task_id'] == task_id:
                    global_pos = idx + 1
                    break
            status = f"🕒 В очереди"
            queue_pos_str = f" (поз. {global_pos})" if global_pos != -1 else ""


        # Имя задачи (укороченное)
        task_name = os.path.basename(task['file_list'][0])
        if len(task_name) > 25: # Обрезаем длинные имена
            task_name = task_name[:22] + "..."
        if len(task['file_list']) > 1:
            task_name += f" (+{len(task['file_list'])-1})"

        tasks_info.append(f"{position_in_user_list}. Задача #{task_id}: {task_name} - {status}{queue_pos_str}")

        # Добавляем кнопку отмены для этой задачи (2 кнопки в ряд)
        builder.add(InlineKeyboardButton(
            text=f"❌ Отменить #{task_id}",
            callback_data=f"cancel:{task_id}:{page}" # Добавляем номер страницы
        ))

    # Подгоняем кнопки отмены по 2 в ряд
    builder.adjust(2)

    # Добавляем кнопки пагинации
    total_pages = (len(user_tasks) + ITEMS_PER_PAGE - 1) // ITEMS_PER_PAGE
    pagination_buttons = []
    if page > 1:
        pagination_buttons.append(InlineKeyboardButton(text="⬅️ Назад", callback_data=f"page:{page-1}"))
    if page < total_pages:
        pagination_buttons.append(InlineKeyboardButton(text="Вперед ➡️", callback_data=f"page:{page+1}"))

    if pagination_buttons:
         builder.row(*pagination_buttons) # Добавляем ряд кнопок пагинации

    text = f"Ваши задачи (Страница {page}/{total_pages}):\n\n" + "\n".join(tasks_info)
    return text, builder.as_markup()

# --- Обработчик кнопок пагинации ---
@router.callback_query(F.data.startswith("page:"))
async def handle_pagination_callback(callback_query: CallbackQuery):
    user_id = callback_query.from_user.id
    try:
        page = int(callback_query.data.split(":")[1])
    except (IndexError, ValueError):
        await callback_query.answer("Ошибка пагинации", show_alert=True)
        return

    message = callback_query.message
    if not message: # Если исходное сообщение было удалено
        await callback_query.answer("Не могу обновить статус.", show_alert=True)
        return

    text, keyboard = await build_task_status(user_id, page)

    if text:
        try:
            await message.edit_text(text, reply_markup=keyboard)
            # Сохраняем новый message_id, если он изменился (обычно нет, но на всякий)
            user_status_messages[user_id] = message.message_id
        except TelegramBadRequest as e:
            # Если сообщение не изменилось или другая ошибка
            print(f"Ошибка редактирования для пагинации: {e}")
            await callback_query.answer("Не удалось обновить статус.") # Краткий ответ пользователю
        except Exception as e:
             print(f"Непредвиденная ошибка пагинации: {e}")
             await callback_query.answer("Произошла ошибка.")
    else:
         # Если задач нет, удаляем сообщение статуса
         await message.delete()
         if user_id in user_status_messages:
             del user_status_messages[user_id]

    await callback_query.answer() # Подтверждаем получение колбэка

@router.message(Command("queue_status"))
async def queue_status(message: Message):
    """
    Проверка статуса очереди с пагинацией.
    """
    if await check_sender(message):
        await message.delete() # Удаляем команду не-пользователя
        return

    user_id = message.from_user.id
    text, keyboard = await build_task_status(user_id, page=1) # Начинаем с первой страницы

    try:
        sent_message = await message.answer(text, reply_markup=keyboard)
        # Сохраняем ID отправленного сообщения для будущих обновлений
        user_status_messages[user_id] = sent_message.message_id
    except Exception as e:
        print(f"Ошибка отправки статуса очереди: {e}")
        # Уведомление об ошибке не отправляем, т.к. основная команда не выполнена

    await message.delete() # Удаляем команду пользователя /queue_status

@router.message(Command("cancel"))
async def cancel_collecting(message: Message, state: FSMContext):
    """
    Отмена текущего сбора файлов.
    """
    if await check_sender(message):
        await message.delete()
        return

    user_id = message.from_user.id
    chat_id = message.chat.id
    current_state = await state.get_state()

    if current_state != MergeStates.collecting.state:
        await send_then_delete(chat_id, "Сбор файлов не был запущен.", delay=5)
        await message.delete()
        return

    # Получаем данные из состояния
    user_data = await state.get_data()
    if user_data.get("_initiator_user_id") != user_id:
        # Пытаются отменить чужой сбор (маловероятно, но проверяем)
        await send_then_delete(chat_id, "Вы не можете отменить сбор, запущенный другим пользователем.", delay=5)
        await message.delete()
        return

    file_list_data = user_data.get('file_list', [])
    saved_message_ids = user_data.get('saved_message_ids', [])
    timeout_task = user_data.get('timeout_task')
    start_message_id = user_data.get('start_message_id')

    # 1. Отменяем таймер автоотмены
    if timeout_task and not timeout_task.done():
        timeout_task.cancel()

    # 2. Удаляем временные файлы
    deleted_files_count = 0
    for file_item in file_list_data:
        file = file_item[0]
        if os.path.exists(file):
            try:
                os.remove(file)
                deleted_files_count += 1
            except OSError as e:
                print(f"Ошибка удаления файла {file} при отмене: {e}")

    # 3. Удаляем сообщения "Файл сохранен"
    deleted_msgs_count = 0
    for msg_id in saved_message_ids:
        try:
            await bot.delete_message(chat_id, msg_id)
            deleted_msgs_count += 1
            await asyncio.sleep(0.1) # Пауза
        except TelegramBadRequest: pass
        except Exception as e: print(f"Ошибка удаления saved_msg {msg_id} при отмене: {e}")

    # 4. Удаляем сообщение /start_merge
    if start_message_id:
        try:
            await bot.delete_message(chat_id, start_message_id)
        except TelegramBadRequest: pass
        except Exception as e: print(f"Ошибка удаления start_msg {start_message_id} при отмене: {e}")

    # 5. Сбрасываем состояние
    await state.clear()

    # 6. Отправляем подтверждение и удаляем команду /cancel
    await send_then_delete(chat_id, f"❌ Сбор файлов отменен. Удалено {deleted_files_count} временных файлов и {deleted_msgs_count} уведомлений.", delay=7)
    await message.delete()

# --- Обработчик отмены конкретной задачи из очереди ---
@router.callback_query(F.data.startswith("cancel:"))
async def handle_cancel_task_callback(callback_query: CallbackQuery):
    user_id = callback_query.from_user.id
    message = callback_query.message # Сообщение со статусом/кнопками

    if not message:
        await callback_query.answer("Не могу отменить задачу (сообщение не найдено).", show_alert=True)
        return

    try:
        parts = callback_query.data.split(":")
        task_id = int(parts[1])
        page = int(parts[2]) if len(parts) > 2 else 1 # Получаем страницу из колбэка
    except (IndexError, ValueError):
        await callback_query.answer("Ошибка данных отмены.", show_alert=True)
        return

    # Ищем задачу в очереди
    task_to_remove = None
    for task in task_queue.queue:
        if task['task_id'] == task_id:
            task_to_remove = task
            break

    if task_to_remove:
        if task_to_remove['user_id'] == user_id:
            # Удаляем задачу из очереди
            task_queue.queue.remove(task_to_remove)
            # Удаляем ее временные файлы
            removed_files_count = 0
            for file in task_to_remove['file_list']:
                if os.path.exists(file):
                    try:
                        os.remove(file)
                        removed_files_count += 1
                    except OSError as e:
                        print(f"Ошибка удаления файла {file} задачи {task_id}: {e}")

            # Обновляем сообщение со статусом
            text, keyboard = await build_task_status(user_id, page=page) # Обновляем ТУ ЖЕ страницу
            if text:
                 try:
                     await message.edit_text(text, reply_markup=keyboard)
                     user_status_messages[user_id] = message.message_id # Обновляем ID на всякий случай
                 except TelegramBadRequest as e:
                     print(f"Не удалось обновить статус после отмены задачи {task_id}: {e}")
                     # Попробуем просто удалить сообщение, если редактирование не удалось
                     await message.delete()
                     if user_id in user_status_messages: del user_status_messages[user_id]

            else: # Если это была последняя задача
                 await message.delete()
                 if user_id in user_status_messages: del user_status_messages[user_id]

            await callback_query.answer(f"Задача #{task_id} удалена из очереди. Файлов удалено: {removed_files_count}")
        else:
            await callback_query.answer("Вы не можете отменить чужую задачу.", show_alert=True)
    else:
        # Проверяем, не выполняется ли задача в данный момент
        if task_id in task_queue.active_tasks:
            if task_queue.active_tasks[task_id]['user_id'] == user_id:
                await callback_query.answer(f"Задача #{task_id} уже выполняется и не может быть отменена.", show_alert=True)
            else:
                 await callback_query.answer("Вы не можете отменить чужую выполняющуюся задачу.", show_alert=True)
        else:
            await callback_query.answer(f"Задача #{task_id} не найдена в очереди или уже завершена.", show_alert=True)
            # Обновляем статус на всякий случай, если задача исчезла
            text, keyboard = await build_task_status(user_id, page=page)
            if text:
                 try:
                     await message.edit_text(text, reply_markup=keyboard)
                     user_status_messages[user_id] = message.message_id
                 except Exception as e:
                     print(f"Не удалось обновить статус (задача #{task_id} не найдена): {e}")
            else:
                 await message.delete()
                 if user_id in user_status_messages: del user_status_messages[user_id]


@router.message(Command("end_merge"))
async def end_merge(message: Message, state: FSMContext):
    """
    Команда для завершения сбора файлов и запроса имени выходного файла.
    """
    if await check_sender(message):
        await message.delete()
        return

    user_id = message.from_user.id
    chat_id = message.chat.id
    current_state = await state.get_state()

    if current_state != MergeStates.collecting.state:
        await send_then_delete(chat_id, "Сбор файлов не был запущен. Введите /start_merge для начала.", delay=5)
        await message.delete()
        return

    # Проверяем инициатора
    user_data = await state.get_data()
    if user_data.get("_initiator_user_id") != user_id:
        await send_then_delete(chat_id, "Завершить сбор может только тот, кто его начал.", delay=5)
        await message.delete()
        return

    file_list_data = user_data.get('file_list', [])
    saved_message_ids = user_data.get('saved_message_ids', [])
    timeout_task = user_data.get('timeout_task')
    start_message_id = user_data.get('start_message_id') # ID сообщения /start_merge

    if not file_list_data:
        await send_then_delete(chat_id, "Нет файлов для обработки! Сбор отменен.", delay=5)
        # Отменяем таймер (если еще работает)
        if timeout_task and not timeout_task.done():
            timeout_task.cancel()
        # Удаляем сообщение /start_merge
        if start_message_id:
             try: await bot.delete_message(chat_id, start_message_id)
             except Exception: pass
        await state.clear()
        await message.delete()
        return

    # 1. Отменяем таймер автоотмены
    if timeout_task and not timeout_task.done():
        timeout_task.cancel()

    # 2. Удаляем сообщения "Файл сохранен"
    deleted_msgs_count = 0
    for msg_id in saved_message_ids:
        try:
            await bot.delete_message(chat_id, msg_id)
            deleted_msgs_count +=1
            await asyncio.sleep(0.1)
        except TelegramBadRequest: pass
        except Exception as e: print(f"Ошибка удаления saved_msg {msg_id} при end_merge: {e}")
    print(f"Удалено {deleted_msgs_count} уведомлений о сохранении файлов.")
    # Очищаем список ID из состояния, чтобы не удалить их снова при ошибке или отмене на следующем шаге
    await state.update_data(saved_message_ids=[])

    # 3. Удаляем сообщение /start_merge
    if start_message_id:
        try:
            await bot.delete_message(chat_id, start_message_id)
        except TelegramBadRequest: pass
        except Exception as e: print(f"Ошибка удаления start_msg {start_message_id} при end_merge: {e}")

    # 4. Переходим к состоянию запроса имени файла
    await state.set_state(MergeStates.naming_file)

    # Создаем клавиатуру с кнопкой "Пропустить"
    keyboard = ReplyKeyboardBuilder()
    keyboard.add(types.KeyboardButton(text="Пропустить"))
    keyboard.adjust(1)

    # Сохраняем ID сообщения /end_merge, чтобы удалить его потом
    await state.update_data(end_merge_message_id=message.message_id)

    await message.answer(
        f"Собрано файлов: {len(file_list_data)}. Введите название для итогового файла (без .docx) "
        "или нажмите 'Пропустить' для использования стандартного имени (merged.docx):",
        reply_markup=keyboard.as_markup(resize_keyboard=True, one_time_keyboard=True) # Скрываем после нажатия
    )
    # Команду /end_merge удалим после получения имени файла

@router.message(MergeStates.naming_file)
async def process_filename(message: Message, state: FSMContext):
    """
    Обработка введенного имени файла или кнопки "Пропустить".
    """
    # Проверяем, что сообщение от того же пользователя, который инициировал
    user_data = await state.get_data()
    if user_data.get("_initiator_user_id") != message.from_user.id:
        # Игнорируем сообщение от другого пользователя
        return

    user_id = message.from_user.id
    chat_id = message.chat.id
    message_thread_id = message.message_thread_id
    is_forum = message.is_topic_message
    file_list_data = user_data.get('file_list', [])

    # Удаляем сообщение с запросом имени и ReplyKeyboard
    # ReplyKeyboard удаляется отправкой нового сообщения с ReplyKeyboardRemove()
    # Сообщение бота с запросом имени нужно удалить явно, если есть его ID
    # Но проще просто отправить новое сообщение с ReplyKeyboardRemove
    await message.answer("Имя принято, формирую задачу...", reply_markup=ReplyKeyboardRemove())

    # Удаляем сообщение /end_merge
    end_merge_message_id = user_data.get('end_merge_message_id')
    if end_merge_message_id:
        try:
            await bot.delete_message(chat_id, end_merge_message_id)
        except Exception as e: print(f"Не удалось удалить сообщение /end_merge: {e}")

    # Удаляем сообщение пользователя с именем файла или "Пропустить"
    await message.delete()

    # Сортируем файлы по ID сообщения (второй элемент кортежа)
    file_list_data.sort(key=lambda x: x[1])

    # Извлекаем только имена файлов после сортировки
    sorted_files = [file[0] for file in file_list_data]

    # Определяем имя выходного файла
    output_file_name_base = "merged" # Имя по умолчанию
    if message.text and message.text.lower() != "пропустить":
        # Очищаем имя файла от недопустимых символов
        raw_name = message.text
        # Удаляем расширение .docx если пользователь его ввел
        if raw_name.lower().endswith(".docx"):
             raw_name = raw_name[:-5]
        # Удаляем недопустимые символы для имен файлов Windows/Linux
        safe_name = re.sub(r'[\\/*?:"<>|]', '', raw_name).strip()
        # Заменяем множественные пробелы на один
        safe_name = re.sub(r'\s+', ' ', safe_name)
        # Ограничиваем длину имени
        safe_name = safe_name[:100] # Ограничим длину базы имени
        if safe_name: # Если после очистки что-то осталось
             output_file_name_base = safe_name
        else: # Если имя стало пустым после очистки
             output_file_name_base = f"file_{int(time.time())}" # Генерируем имя

    output_file_name = output_file_name_base + ".docx"

    # Добавляем задачу в очередь с отсортированным списком файлов
    queue_position = task_queue.add_task(user_id, chat_id, message_thread_id, is_forum, sorted_files, output_file_name)

    # --- НЕ ИСПОЛЬЗУЕМ ReplyKeyboard здесь ---
    # # Возвращаем обычную клавиатуру
    # keyboard = ReplyKeyboardBuilder()
    # keyboard.add(types.KeyboardButton(text="/start_merge"))
    # keyboard.add(types.KeyboardButton(text="/end_merge")) # Не нужен, т.к. мы уже закончили
    # keyboard.add(types.KeyboardButton(text="/cancel")) # Не нужен
    # keyboard.add(types.KeyboardButton(text="/queue_status"))
    # keyboard.add(types.KeyboardButton(text="/limits"))
    # keyboard.adjust(2)

    confirmation_text = (
        f"✅ Задача на объединение создана.\n"
        f"Итоговый файл будет назван: `{output_file_name}`\n"
    )
    if queue_position > 1:
        # Оцениваем время ожидания очень грубо (зависит от длины активных задач)
        wait_estimate = len(task_queue.queue) # Примерно равно позиции
        confirmation_text += (f"Ваша задача добавлена в очередь (позиция {queue_position}).\n"
                              #f"Примерное время ожидания: {wait_estimate * 1} минут(ы).\n" # Убрал время, оно неточное
                              f"Используйте /queue_status для проверки статуса.")
    else:
        confirmation_text += "Ваша задача будет обработана в ближайшее время."

    await bot.send_message(chat_id, confirmation_text, parse_mode="Markdown")

    # Очищаем состояние FSM после успешного добавления задачи
    await state.clear()

    # Пытаемся запустить обработку задачи, если есть свободные потоки
    asyncio.create_task(check_and_process_queue())

async def check_and_process_queue():
    """
    Проверяет очередь и запускает обработку новых задач, если есть свободные ресурсы.
    """
    while task_queue.can_process_now():
        task = task_queue.get_next_task()
        if task:
            user_id = task['user_id']
            chat_id = task['chat_id']
            message_thread_id = task['message_thread_id']
            is_forum = task['is_forum']
            file_list = task['file_list']
            output_file_name = task['output_file_name']
            task_id = task['task_id']

            send_kwargs = {} # Словарь для тем
            if is_forum and message_thread_id: # Проверяем, существуют ли темы в группе
                    send_kwargs["message_thread_id"] = message_thread_id

            # Обновляем сообщение со статусом пользователя, если оно есть
            if user_id in user_status_messages:
                 message_id = user_status_messages[user_id]
                 try:
                     # Получаем текущую страницу из старой клавиатуры
                     current_page = 1
                     try:
                         msg = await bot.edit_message_reply_markup(chat_id, message_id, reply_markup=None)
                         if msg and msg.reply_markup:
                            for row in msg.reply_markup.inline_keyboard:
                                 for button in row:
                                     if button.callback_data.startswith("page:"):
                                         current_page = int(button.callback_data.split(":")[1])
                                         break
                                 if current_page > 1: break
                     except Exception: pass

                     text, keyboard = await build_task_status(user_id, page=current_page) # Перестраиваем для текущей страницы
                     if text:
                          await bot.edit_message_text(text, chat_id, message_id, reply_markup=keyboard)
                     else: # Задач не осталось
                         await bot.delete_message(chat_id, message_id)
                         del user_status_messages[user_id]
                 except Exception as e:
                     print(f"Ошибка обновления статуса перед началом задачи {task_id}: {e}")
                     # Не страшно, продолжаем выполнение

            # Уведомляем пользователя о начале обработки (сообщение будет висеть до результата)
            processing_msg = await bot.send_message(
                chat_id,
                f"⚙️ Начинаю обработку задачи #{task_id} ({os.path.basename(output_file_name)}, {len(file_list)} файлов). Это может занять некоторое время...",
                **send_kwargs
                )

            # Запускаем обработку в фоновом режиме
            asyncio.create_task(process_and_merge_files_with_queue(
                chat_id, send_kwargs, file_list, output_file_name, task_id, user_id, processing_msg.message_id
                ))

async def process_and_merge_files_with_queue(chat_id, send_kwargs, file_list, output_file_name, task_id, user_id, processing_msg_id):
    """
    Асинхронная функция для обработки и объединения файлов с учетом очереди.
    """
    start_time = time.time()
    result_file_path = None
    error_occurred = False
    error_message = ""
    original_files_to_delete = list(file_list) # Копируем список для удаления
    converted_files_to_delete = []

    try:
        # 1. Конвертация файлов (если требуется)
        print(f"[Task {task_id}] Начало конвертации {len(file_list)} файлов...")
        converted_files = await process_files(file_list)
        # Сохраняем имена сконвертированных файлов для последующего удаления
        # Исключаем те, что были исходными docx
        original_docx = {f for f in file_list if f.lower().endswith('.docx')}
        converted_files_to_delete = [f for f in converted_files if f not in original_docx and os.path.exists(f)]
        print(f"[Task {task_id}] Конвертация завершена. Получено {len(converted_files)} файлов для слияния.")

        if not converted_files:
             raise ValueError("Нет файлов для объединения после этапа конвертации.")

        # 2. Объединение файлов
        print(f"[Task {task_id}] Начало слияния {len(converted_files)} файлов в {output_file_name}...")
        merged_file = await merge_docx(converted_files, output_file_name)
        result_file_path = merged_file # Сохраняем путь к результату
        print(f"[Task {task_id}] Слияние завершено.")

        # 3. Отправка результата
        if os.path.exists(merged_file) and os.path.getsize(merged_file) > 0:
             # Формируем сообщение с информацией
             duration = time.time() - start_time
             file_list_str = "\n".join([f"- `{os.path.basename(f)}`" for f in file_list]) # Показываем исходные имена
             success_message = (
                 f"✅ Задача #{task_id} успешно завершена за {duration:.1f} сек!\n"
                 f"Файл: `{os.path.basename(output_file_name)}`\n"
                 f"Объединено из {len(file_list)} файлов:\n{file_list_str}"
             )
             await bot.send_message(chat_id, success_message, parse_mode="Markdown", **send_kwargs)

             # Отправляем сам документ
             document = FSInputFile(merged_file)
             await bot.send_document(chat_id, document=document, caption=f"Результат задачи #{task_id}", **send_kwargs)
        else:
             # Если файл пустой или не существует после слияния
             raise ValueError(f"Итоговый файл '{output_file_name}' пуст или не найден после слияния.")

    except Exception as e:
        error_occurred = True
        error_message = str(e)
        print(f"[Task {task_id}] Ошибка обработки: {e}")
        # Отправляем сообщение об ошибке
        try:
            await bot.send_message(chat_id, f"❌ Произошла ошибка при обработке задачи #{task_id}: {e}", **send_kwargs)
        except Exception as send_err:
            print(f"Не удалось отправить сообщение об ошибке для задачи {task_id}: {send_err}")
    finally:
        # 4. Удаляем сообщение "Начинаю обработку..."
        try:
            await bot.delete_message(chat_id, processing_msg_id)
        except Exception as del_err:
            print(f"Не удалось удалить сообщение о начале обработки {processing_msg_id}: {del_err}")

        # 5. Удаляем временные файлы (исходные и сконвертированные)
        files_to_clean = set(original_files_to_delete + converted_files_to_delete)
        if result_file_path and os.path.exists(result_file_path):
            files_to_clean.add(result_file_path) # Добавляем итоговый файл для удаления

        print(f"[Task {task_id}] Очистка временных файлов: {len(files_to_clean)} шт.")
        for file_path in files_to_clean:
            if os.path.exists(file_path):
                try:
                    os.remove(file_path)
                except OSError as rm_err:
                    print(f"Ошибка удаления временного файла {file_path}: {rm_err}")

        # 6. Отмечаем задачу как выполненную (и обновляем статус)
        await task_queue.complete_task(task_id) # Передаем task_id

        # 7. Проверяем, можно ли обработать следующую задачу
        asyncio.create_task(check_and_process_queue())


@router.message(F.document)
async def handle_document(message: Message, state: FSMContext):
    """
    Обработчик полученных файлов.
    Если сбор файлов запущен, сохраняет полученный документ на диск
    и добавляет его имя в список для дальнейшей обработки.
    """
    # Сначала проверяем отправителя, чтобы не выполнять лишних действий
    if await check_sender(message):
        await message.delete() # Удаляем сообщение от не-пользователя
        return

    # Проверяем состояние FSM
    current_state = await state.get_state()
    if current_state != MergeStates.collecting.state:
        await send_then_delete(message.chat.id, "Сбор файлов не запущен. Введите /start_merge для начала.", delay=5)
        # Не удаляем сам файл, т.к. пользователь мог его случайно прислать
        return

    # Проверяем, что файл прислал тот же пользователь, кто начал сбор
    user_data = await state.get_data()
    if user_data.get("_initiator_user_id") != message.from_user.id:
        await send_then_delete(message.chat.id, "Сейчас идет сбор файлов другим пользователем. Вы не можете добавлять файлы.", delay=7)
        return

    user_id = message.from_user.id
    chat_id = message.chat.id
    doc = message.document

    # Проверяем формат файла ДО скачивания
    file_name_orig = doc.file_name if doc.file_name else "unknown_file"
    extension = os.path.splitext(file_name_orig)[1].lower()
    allowed_extensions = (".docx", ".fb2", ".txt", ".epub")

    if extension not in allowed_extensions:
        await send_then_delete(chat_id,
            f"⚠️ Неподдерживаемый формат файла: `{file_name_orig}`. "
            f"Допустимые форматы: {', '.join(allowed_extensions)}",
            delay=7, parse_mode="Markdown")
        return # Не скачиваем и не сохраняем

    file_size = doc.file_size if doc.file_size else 0

    # Проверяем лимиты пользователя (размер и количество)
    lock = user_limits.get_lock(user_id)
    async with lock:
        is_allowed, error_msg = user_limits.check_limits(user_id, file_size)
        if not is_allowed:
            await send_then_delete(chat_id, error_msg, delay=7)
            return # Выходим, лимит превышен

        # Если лимит позволяет, СРАЗУ увеличиваем счетчик ВНУТРИ блокировки
        user_limits.increment_counter(user_id)
        max_files = user_limits.max_files
        files_today_count = user_limits.user_data[user_id]['files_today']
        # Блокировка освобождается здесь

    # --- Операции вне блокировки (загрузка, сохранение) ---
    # Создаем уникальное имя файла для сохранения на диске
    # Добавляем user_id и timestamp для большей уникальности
    timestamp = int(time.time())
    safe_base_name = re.sub(r'[\\/*?:"<>|]', '', os.path.splitext(file_name_orig)[0])[:50] # Очищаем и обрезаем имя
    unique_file_name = f"user_{user_id}_{timestamp}_{safe_base_name}{extension}"
    # Убираем возможные двойные точки перед расширением
    unique_file_name = unique_file_name.replace("..", ".")

    # Создаем директорию 'temp_files', если ее нет
    temp_dir = "temp_files"
    os.makedirs(temp_dir, exist_ok=True)
    save_path = os.path.join(temp_dir, unique_file_name)

    try:
        # Скачиваем файл
        print(f"Скачивание файла {file_name_orig} от user {user_id}...")
        file_info = await bot.get_file(doc.file_id)
        # Используем download для получения байтов
        downloaded_file_bytes = await bot.download_file(file_info.file_path)

        # Сохраняем файл на диск асинхронно
        async with aiofiles.open(save_path, 'wb') as new_file:
            await new_file.write(downloaded_file_bytes.read())
        print(f"Файл сохранен как: {save_path}")

        # Добавляем файл в список состояния FSM
        user_data = await state.get_data() # Получаем актуальные данные
        file_list = user_data.get('file_list', [])
        # Храним кортеж (путь_к_файлу, id_сообщения_с_файлом)
        file_list.append((save_path, message.message_id))

        # Получаем ID сообщений "Файл сохранен"
        saved_message_ids = user_data.get('saved_message_ids', [])

        # Отправляем подтверждение и сохраняем его ID
        confirmation_msg = await message.answer(
            f"✅ Файл `{file_name_orig}` ({len(file_list)} шт.) добавлен.\n"
            f"Лимит: {files_today_count}/{max_files}",
            parse_mode="Markdown"
        )
        saved_message_ids.append(confirmation_msg.message_id)

        # Обновляем состояние с новым списком файлов и ID сообщения
        await state.update_data(file_list=file_list, saved_message_ids=saved_message_ids)

        # Перезапускаем таймер автоотмены (сбрасываем таймаут при активности)
        timeout_task = user_data.get('timeout_task')
        if timeout_task and not timeout_task.done():
            timeout_task.cancel() # Отменяем старый таймер
        # Запускаем новый таймер
        start_msg_id = user_data.get('start_message_id')
        new_timeout_handle = asyncio.create_task(
             collection_timeout_task(chat_id, user_id, state, start_msg_id)
        )
        await state.update_data(timeout_task=new_timeout_handle)
        print(f"Таймер автоотмены перезапущен для user {user_id}")


    except Exception as e:
        print(f"Ошибка при обработке документа от user {user_id}: {e}")
        await send_then_delete(chat_id, f"❌ Ошибка при сохранении файла `{file_name_orig}`: {e}", delay=7, parse_mode="Markdown")
        # Важно: откатываем счетчик лимита, если файл не был сохранен
        async with lock:
             if user_id in user_limits.user_data and user_limits.user_data[user_id]['files_today'] > 0:
                 user_limits.user_data[user_id]['files_today'] -= 1
                 print(f"Счетчик лимита откачен для user {user_id} из-за ошибки сохранения.")
    # Не удаляем сообщение с файлом пользователя

# ===================== Прочие команды =====================
@router.message(Command("start"))
async def send_welcome(message: Message):
    if await check_sender(message):
        await message.delete()
        return
    keyboard = ReplyKeyboardBuilder()
    keyboard.add(types.KeyboardButton(text="/info"))
    keyboard.add(types.KeyboardButton(text="/start_merge"))
    keyboard.adjust(2)
    await message.answer("👋 Привет! Я бот для объединения файлов.\n"
                         "Нажмите /info для получения справки или /start_merge для начала.",
                         reply_markup=keyboard.as_markup(resize_keyboard=True))
    await message.delete()

@router.message(Command("info"))
async def send_info(message: Message):
    if await check_sender(message):
        await message.delete()
        return

    keyboard = ReplyKeyboardBuilder()
    keyboard.add(types.KeyboardButton(text="/start_merge"))
    keyboard.add(types.KeyboardButton(text="/limits"))
    keyboard.add(types.KeyboardButton(text="/queue_status"))
    keyboard.add(types.KeyboardButton(text="/cancel")) # Добавим для удобства
    keyboard.adjust(2)

    max_files = user_limits.max_files
    max_size = user_limits.max_size

    await message.answer(
        "ℹ️ *Информация о боте*\n\n"
        "Я могу объединять несколько файлов форматов `.docx`, `.fb2`, `.epub`, `.txt` в один итоговый `.docx` файл.\n\n"
        "⚙️ *Как использовать:*\n"
        "1. Нажмите `/start_merge`, чтобы начать сбор.\n"
        "2. Отправьте мне файлы по одному.\n"
        "3. Когда все файлы отправлены, нажмите `/end_merge`.\n"
        "4. Введите желаемое имя для итогового файла или нажмите 'Пропустить'.\n"
        "5. Дождитесь обработки и получения результата.\n\n"
        "*Важно:* Файлы объединяются в том порядке, в котором вы их отправили.\n\n"
        "📊 *Лимиты:*\n"
        f"• *Количество:* {max_files} файлов в сутки на пользователя.\n"
        f"• *Размер:* Максимум {max_size} МБ на один файл.\n"
        f"• *Сброс лимитов:* Ежедневно в 00:00 по UTC.\n\n"
        "⏱ *Автоотмена:* Сбор файлов будет автоматически отменен, если вы не будете отправлять файлы или не нажмете /end_merge в течение + f"{int(COLLECTION_TIMEOUT/60)} минут.\n\n"
        "*Команды:*\n"
        "`/start_merge` - Начать новый сбор файлов.\n"
        "`/end_merge` - Завершить сбор и запустить объединение.\n"
        "/cancel - Отменить текущий сбор файлов.\n"
        "`/limits` - Показать текущее использование лимитов.\n"
        "`/queue_status` - Показать статус ваших задач в очереди.\n"
        "/info - Показать это сообщение.",
        parse_mode="Markdown",
        reply_markup=keyboard.as_markup(resize_keyboard=True)
    )
    await message.delete()

@router.message(Command("limits"))
async def check_limits_cmd(message: Message): # Переименовал, чтобы не конфликтовать с функцией check_limits
    """Показывает текущие лимиты и время до сброса."""
    if await check_sender(message):
        await message.delete()
        return

    user_id = message.from_user.id

    # Принудительно проверяем сброс лимитов перед показом
    now = datetime.now(timezone.utc)
    if now >= user_limits.last_global_reset + timedelta(days=1):
        print("Сброс суточных лимитов (проверка перед /limits)...")
        user_limits.user_data.clear()
        user_limits.last_global_reset = user_limits._get_last_utc_midnight()

    next_reset = user_limits.last_global_reset + timedelta(days=1)
    time_left = next_reset - now
    # Используем total_seconds() для корректного расчета оставшегося времени
    total_seconds_left = max(0, time_left.total_seconds()) # Убедимся, что не отрицательное
    hours_left = int(total_seconds_left // 3600)
    minutes_left = int((total_seconds_left % 3600) // 60)

    max_files = user_limits.max_files
    max_size = user_limits.max_size

    # Получаем текущее использование (может быть 0, если записи нет)
    files_used = user_limits.user_data.get(user_id, {}).get('files_today', 0)
    files_left = max(0, max_files - files_used) # Не может быть меньше 0

    await message.answer(
        f"📊 *Ваши лимиты:*\n\n"
        f"• Использовано файлов сегодня: *{files_used}* из *{max_files}*\n"
        f"• Осталось файлов на сегодня: *{files_left}*\n"
        f"• Максимальный размер 1 файла: *{max_size} MB*\n\n"
        f"🔄 Лимиты сбрасываются в *00:00 UTC* (примерно через *{hours_left} ч. {minutes_left} мин.*)",
        parse_mode="Markdown"
    )
    await message.delete()

# ===================== Запуск бота =====================
async def main():
    # Создаем директорию для временных файлов, если ее нет
    if not os.path.exists("temp_files"):
        os.makedirs("temp_files")

    # Настройка хранения состояний FSM в памяти
    storage = MemoryStorage()
    # Создание диспетчера и подключение роутера
    dp = Dispatcher(storage=storage)
    dp.include_router(router)

    # Удаление вебхука перед запуском поллинга (если он был установлен)
    await bot.delete_webhook(drop_pending_updates=True)

    print("Бот запускается...")
    # Запуск поллинга с разрешенными обновлениями
    # allowed_updates можно уточнить для оптимизации, но для простоты оставим по умолчанию
    await dp.start_polling(bot, allowed_updates=dp.resolve_used_update_types())
    print("Бот остановлен.")

if __name__ == "__main__":
    try:
        asyncio.run(main())
    except KeyboardInterrupt:
        print("Завершение работы бота...")
    finally:
        # Закрываем пул потоков при завершении
        thread_pool.shutdown(wait=True)
        print("Пул потоков остановлен.")
