#!/usr/bin/env python3
# -*- coding: utf-8 -*-

# Установка зависимостей (если запускаете в Jupyter Notebook или аналогичном окружении)
import os
import re
import time
import docx
from docx import Document
from docxcompose.composer import Composer
from bs4 import BeautifulSoup
import ebooklib
from ebooklib import epub
from aiogram import Bot, Router, types, F, Dispatcher
from aiogram.types import Message
from aiogram.types import FSInputFile
from aiogram.filters import Command
from aiogram.utils.keyboard import ReplyKeyboardBuilder
from aiogram.utils import markdown as md
import aiofiles
import asyncio
import nest_asyncio
nest_asyncio.apply()

# Декоратор для измерения времени выполнения функции
def timer(func):
    async def wrapper(*args, **kwargs):
        start_time = time.time()
        result = await func(*args, **kwargs)
        elapsed = time.time() - start_time
        print(f"[PROFILING] Функция {func.__name__} выполнилась за {elapsed:.2f} секунд")
        return result
    return wrapper

# Замените токен на свой
API_TOKEN = '7692853253:AAHGfbJhall58TafIqTBdAujVnuXhhHCwYk'

bot = Bot(token=API_TOKEN)
router = Router()

# Глобальный словарь для хранения состояния пользователей
user_states = {}

# Имя итогового файла по умолчанию
output_file_name = "merged.docx"

# ===================== Функции конвертации =====================
@timer
async def convert_epub_to_docx(epub_file, docx_file):
    # Открываем EPUB-файл
    book = epub.read_epub(epub_file)
    document = Document()
    # Перебираем элементы книги
    for item in book.get_items():
        if item.get_type() == ebooklib.ITEM_DOCUMENT:
            soup = BeautifulSoup(item.content, 'html.parser')
            for element in soup.find_all():
                if element.name == 'h1':
                    document.add_heading(element.get_text(), level=0)
                elif element.name == 'p':
                    doc_paragraph = document.add_paragraph()
                    # Перебор вложенных элементов абзаца
                    for sub in element.contents:
                        if hasattr(sub, 'name'):
                            if sub.name == 'strong':
                                run = doc_paragraph.add_run(sub.get_text())
                                run.bold = True
                            elif sub.name == 'em':
                                run = doc_paragraph.add_run(sub.get_text())
                                run.italic = True
                        else:
                            # Если это просто текст
                            doc_paragraph.add_run(sub)
    document.save(docx_file)

@timer
async def convert_fb2_to_docx(fb2_file, docx_file):
    async with aiofiles.open(fb2_file, 'r', encoding='utf-8') as f:
        content = await f.read()
    soup = BeautifulSoup(content, 'xml')
    document = Document()
    for element in soup.find_all():
        if element.name == 'title':
            document.add_heading(element.get_text(), level=0)
        elif element.name == 'p':
            # Если абзац не является частью title или annotation
            if element.find_parent(['title', 'annotation']) is None:
                doc_paragraph = document.add_paragraph()
                for sub in element.contents:
                    if hasattr(sub, 'name'):
                        if sub.name == 'strong':
                            run = doc_paragraph.add_run(sub.get_text())
                            run.bold = True
                        elif sub.name == 'emphasis':
                            run = doc_paragraph.add_run(sub.get_text())
                            run.italic = True
                    else:
                        doc_paragraph.add_run(sub)
    document.save(docx_file)

@timer
async def convert_txt_to_docx(txt_file, docx_file):
    async with aiofiles.open(txt_file, 'r', encoding='utf-8') as f:
        text = await f.read()
    document = Document()
    for line in text.splitlines():
        document.add_paragraph(line)
    document.save(docx_file)

@timer
async def process_files(file_list):
    """
    Обрабатывает список файлов, конвертируя их в формат .docx (если требуется)
    и возвращает список имен файлов в формате .docx для последующего объединения.
    """
    converted_files = []
    for file in file_list:
        ext = os.path.splitext(file)[1].lower()
        # Если файл уже в формате .docx – добавляем его в список
        if ext == ".docx":
            converted_files.append(file)
        elif ext == ".txt":
            docx_file = os.path.splitext(file)[0] + ".docx"
            await convert_txt_to_docx(file, docx_file)
            converted_files.append(docx_file)
        elif ext == ".fb2":
            docx_file = os.path.splitext(file)[0] + ".docx"
            await convert_fb2_to_docx(file, docx_file)
            converted_files.append(docx_file)
        elif ext == ".epub":
            docx_file = os.path.splitext(file)[0] + ".docx"
            await convert_epub_to_docx(file, docx_file)
            converted_files.append(docx_file)
    return converted_files

# ===================== Функции для работы с документами =====================
@timer
async def check_and_add_title(doc, file_name):
    """
    Проверяет первые абзацы документа на наличие заголовка (например, "Глава ...").
    Если заголовок не найден, добавляет его на основе имени файла.
    """
    patterns = [
        r'Глава[ ]{0,4}\d{1,4}',
        r'Часть[ ]{0,4}\d{1,4}',
        r'^Пролог[ .!]*$',
        r'^Описание[ .!]*$',
        r'^Аннотация[ .!]*$',
        r'^Annotation[ .!]*$',
        r'^Предисловие от автора[ .!]*$'
    ]
    if doc.paragraphs:
        check_paragraphs = doc.paragraphs[0:4]
        title_found = False
        for p in check_paragraphs:
            for pattern in patterns:
                if re.search(pattern, p.text):
                    title_found = True
                    break
            if title_found:
                break
        if not title_found:
            # Добавляем заголовок перед первым абзацем
            title = os.path.splitext(os.path.basename(file_name))[0]
            title_run = doc.paragraphs[0].insert_paragraph_before().add_run(f"{title}\n")
            # Форматирование заголовка
            title_run.bold = True
    return doc

@timer
async def merge_docx(file_list, output_file_name):
    # Создаем новый документ
    merged_document = Document(file_list[0])
    merged_document = await check_and_add_title(merged_document, file_list[0])
    composer = Composer(merged_document)
    for file in file_list[1:]:
        doc = Document(file)
        # Проверяем и добавляем название главы при необходимости
        doc = await check_and_add_title(doc, file)
        composer.append(doc)
    # Сохраняем итоговый документ
    composer.save(output_file_name)
    print(f"Файлы объединены в {output_file_name}")
    return output_file_name

# ===================== Обработчики Telegram-бота =====================
@router.message(Command("start_merge"))
async def start_merge(message: Message):
    """
    Команда для начала сбора файлов.
    """
    chat_id = message.chat.id
    if chat_id in user_states and user_states[chat_id]['is_collecting']:
        await message.answer("Сбор файлов уже запущен.")
        return
    user_states[chat_id] = {'is_collecting': True, 'file_list': []}
    await message.answer("Сбор файлов начат! Отправляйте файлы. Используйте /end_merge для завершения.")

@router.message(Command("end_merge"))
async def end_merge(message: Message):
    """
    Команда для завершения сбора файлов и запуска объединения.
    """
    chat_id = message.chat.id
    if chat_id not in user_states or not user_states[chat_id]['is_collecting']:
        await message.answer("Сбор файлов не был запущен. Введите /start_merge для начала.")
        return
    file_list = user_states[chat_id]['file_list']
    if not file_list:
        await message.answer("Нет файлов для обработки!")
        user_states[chat_id]['is_collecting'] = False
        return
    # Обработка и конвертация файлов
    converted_files = await process_files(file_list)
    merged_file = await merge_docx(converted_files, output_file_name)
    # Формируем сообщение с информацией о собранных файлах
    file_list_str = "\n".join(file_list)
    await message.answer(f"Файлы объединены в {output_file_name}.\nСобрано {len(file_list)} файлов:\n{file_list_str}")
    # Отправляем объединённый файл обратно пользователю
    try:
       document = FSInputFile(merged_file)
       await message.answer_document(document, caption="Ваш объединённый документ")
    except Exception as e:
        await message.answer(f"Ошибка при отправке объединённого файла: {str(e)}")
    # Сброс состояния
    user_states[chat_id] = {'is_collecting': False, 'file_list': []}

@router.message(F.document)
async def handle_document(message: Message):
    """
    Обработчик полученных файлов.
    Если сбор файлов запущен, сохраняет полученный документ на диск
    и добавляет его имя в список для дальнейшей обработки.
    """
    chat_id = message.chat.id
    if chat_id not in user_states or not user_states[chat_id]['is_collecting']:
        await message.answer("Сбор файлов не запущен. Введите /start_merge для начала.")
        return
    try:
        file_info = await bot.get_file(message.document.file_id)
        downloaded_file = await bot.download_file(file_info.file_path)
        file_name = message.document.file_name
        base_name, extension = os.path.splitext(file_name)
        counter = 1
        if extension not in (".docx", ".fb2", ".txt", ".epub"):
            await message.answer(f"Неизвестный формат файла: {file_name}. Пожалуйста, отправляйте файлы только в форматах docx, fb2, epub, txt.")
        else:
            # Проверяем, существует ли файл, и добавляем суффикс, если нужно
            while os.path.exists(file_name):
                file_name = f"{base_name}({counter}){extension}"
                counter += 1
            # Сохраняем файл на диск
            async with aiofiles.open(file_name, 'wb') as new_file:
                await new_file.write(downloaded_file.read())
            user_states[chat_id]['file_list'].append(file_name)
            await message.answer(f"Файл {file_name} сохранён!")
    except Exception as e:
        await message.answer(f"Ошибка при сохранении файла: {str(e)}")

@router.message(Command("start"))
async def send_welcome(message: Message):
    await message.answer("Привет, я бот для объединения файлов! Нажми /info для получения дополнительной информации.")

@router.message(Command("info"))
async def send_info(message: Message):
    await message.answer("Данный бот объединяет файлы в формате docx. Если формат файла не docx, то файл конвертируется в этот формат и объединяется. Для конвертации поддерживаются форматы fb2, epub, txt. В процессе конвертации сохраняется лишь текст, жирный и курсивный формат и оглавление, всё остальное теряется. Для начала работы нажми на /start_merge. После отправь файлы и нажми на /end_merge, чтобы завершить работу. Удачи!")

# ===================== Запуск бота =====================
async def main():
    dp = Dispatcher()
    dp.include_router(router)
    print("Бот запущен.")
    await dp.start_polling(bot)

if __name__ == "__main__":
    asyncio.run(main())

print("Файлы в текущей директории:", os.listdir())
